#!/usr/bin/env python3
"""
Generate Daily Stock Analysis Whitepaper (.docx)
Professional executive-style document using python-docx.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Color Constants ────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x2B, 0x57, 0x9A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_HEX = "F2F2F2"
DARK_BLUE_HEX = "2B579A"
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
ACCENT_BLUE = RGBColor(0x1A, 0x73, 0xE8)
CALLOUT_BORDER_HEX = "2B579A"
CALLOUT_BG_HEX = "EBF0F9"


# ─── Helper Functions ───────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_text(cell, text, bold=False, color=None, size=None, font_name=None, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
    # Reduce cell paragraph spacing
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return run


def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table with dark blue headers and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_shading(cell, DARK_BLUE_HEX)
        set_cell_text(cell, header_text, bold=True, color=WHITE, size=Pt(9.5),
                      font_name="Calibri")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY_HEX)
            set_cell_text(cell, str(cell_text), size=Pt(9), font_name="Calibri")

    # Apply column widths if specified
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width

    return table


def add_callout_box(doc, text):
    """Add a bordered callout box paragraph with accent styling."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)

    # Add border via XML
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{CALLOUT_BORDER_HEX}"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{CALLOUT_BORDER_HEX}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{CALLOUT_BORDER_HEX}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="{CALLOUT_BORDER_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)

    # Add background shading
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{CALLOUT_BG_HEX}" w:val="clear"/>'
    )
    pPr.append(shading)

    # Add key metric icon and text
    run_icon = p.add_run("KEY METRIC  |  ")
    run_icon.bold = True
    run_icon.font.size = Pt(9)
    run_icon.font.color.rgb = DARK_BLUE
    run_icon.font.name = "Calibri"

    run_text = p.add_run(text)
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = DARK_GRAY
    run_text.font.name = "Calibri"
    run_text.bold = True

    return p


def add_body_text(doc, text, bold_terms=None):
    """Add a body paragraph with optional bold terms on first use."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)

    if bold_terms is None:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
    else:
        # Split and bold specific terms
        remaining = text
        for term in bold_terms:
            if term in remaining:
                idx = remaining.index(term)
                if idx > 0:
                    run = p.add_run(remaining[:idx])
                    run.font.size = Pt(10.5)
                    run.font.name = "Calibri"
                    run.font.color.rgb = DARK_GRAY
                run_bold = p.add_run(term)
                run_bold.bold = True
                run_bold.font.size = Pt(10.5)
                run_bold.font.name = "Calibri"
                run_bold.font.color.rgb = DARK_GRAY
                remaining = remaining[idx + len(term):]
        if remaining:
            run = p.add_run(remaining)
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY

    return p


def add_spacer(doc, pt=6):
    """Add a small spacer paragraph."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(pt)
    run = p.add_run("")
    run.font.size = Pt(2)
    return p


def add_section_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri Light"
        run.font.color.rgb = DARK_BLUE
    return h


def create_flow_table(doc, steps, title_col="Step", desc_col="Description", detail_col=None):
    """Create a numbered flow/process table with arrow indicators."""
    headers = ["#", title_col, desc_col]
    if detail_col:
        headers.append(detail_col)

    rows = []
    for i, step in enumerate(steps):
        if detail_col:
            rows.append([str(i + 1), step[0], step[1], step[2]])
        else:
            rows.append([str(i + 1), step[0], step[1]])

    return create_styled_table(doc, headers, rows)


# ─── Document Generation ────────────────────────────────────────────

def generate_whitepaper():
    doc = Document()

    # ─── Page Setup ─────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ─── Set Default Font ───────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = DARK_GRAY

    # Update heading styles
    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.name = "Calibri Light"
        hs.font.color.rgb = DARK_BLUE
        if i == 1:
            hs.font.size = Pt(22)
        elif i == 2:
            hs.font.size = Pt(16)
        elif i == 3:
            hs.font.size = Pt(13)
        else:
            hs.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════

    # Add vertical spacing before title
    for _ in range(6):
        add_spacer(doc, 24)

    # Decorative line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("_" * 60)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(14)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_title.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    run = p_title.add_run("Daily Stock Analysis Pipeline")
    run.bold = True
    run.font.size = Pt(32)
    run.font.name = "Calibri Light"
    run.font.color.rgb = DARK_BLUE

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p_sub.paragraph_format
    pf.space_after = Pt(6)
    run = p_sub.add_run("Technical White Paper")
    run.font.size = Pt(18)
    run.font.name = "Calibri Light"
    run.font.color.rgb = MEDIUM_GRAY

    # Second decorative line
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line2.add_run("_" * 60)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(14)

    # Version / Author / Date block
    add_spacer(doc, 18)

    for label, value in [
        ("Version", "3.0"),
        ("Author", "cdemchalk"),
        ("Date", "February 2026"),
        ("Classification", "Internal — Technology Executive Audience"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        run_l = p.add_run(f"{label}:  ")
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_l.font.name = "Calibri"
        run_l.font.color.rgb = DARK_BLUE
        run_v = p.add_run(value)
        run_v.font.size = Pt(11)
        run_v.font.name = "Calibri"
        run_v.font.color.rgb = DARK_GRAY

    add_spacer(doc, 24)

    # Tag line
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tag.add_run("Serverless Investment Intelligence on Azure Functions")
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    run.font.color.rgb = MEDIUM_GRAY

    # Page break after title page
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "Table of Contents", level=1)

    toc_items = [
        ("1.", "Executive Summary"),
        ("2.", "System Architecture"),
        ("3.", "Data Pipeline"),
        ("4.", "Technical Analysis Engine"),
        ("5.", "Options Analysis & Strategy Engine"),
        ("6.", "Black-Scholes Backtesting Framework"),
        ("7.", "AI Summarization (GPT-5)"),
        ("8.", "Report Generation"),
        ("9.", "Cloud Infrastructure (Azure)"),
        ("10.", "Cost Analysis"),
        ("11.", "Roadmap"),
    ]

    for num, title in toc_items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.left_indent = Cm(1.0)
        run_n = p.add_run(f"{num}  ")
        run_n.bold = True
        run_n.font.size = Pt(11)
        run_n.font.name = "Calibri"
        run_n.font.color.rgb = DARK_BLUE
        run_t = p.add_run(title)
        run_t.font.size = Pt(11)
        run_t.font.name = "Calibri"
        run_t.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "1. Executive Summary", level=1)

    add_body_text(doc,
        "The Daily Stock Analysis Pipeline is a fully automated, serverless system that "
        "transforms raw market data into actionable investment intelligence. Running on "
        "Microsoft Azure Functions with Python 3.12, it executes every weekday at 1:00 PM UTC "
        "and delivers a comprehensive HTML dashboard report via email.",
        bold_terms=["Daily Stock Analysis Pipeline", "Azure Functions"]
    )

    add_body_text(doc,
        "The system analyzes a configurable watchlist of equity tickers across five dimensions: "
        "22 technical indicators, approximately 22 fundamental data points, full options chain "
        "analysis with 7 algorithmically-scored strategies, market sentiment from social platforms, "
        "and AI-generated summaries powered by OpenAI GPT-5. Each run produces a single, "
        "information-dense email report that a portfolio manager can read in under five minutes."
    )

    add_body_text(doc,
        "Three access modes are supported: a scheduled timer trigger for daily automated delivery, "
        "an HTTP API for on-demand analysis, and a command-line interface for local development "
        "and backtesting. The architecture introduces zero new pip dependencies beyond the existing "
        "requirements, keeping the deployment lean and the attack surface small."
    )

    add_spacer(doc, 4)

    # Key metric callouts
    add_callout_box(doc, "22 technical indicators computed per ticker per run")
    add_callout_box(doc, "7 algorithmically-scored options strategies with exact risk profiles")
    add_callout_box(doc, "Black-Scholes backtesting validates strategies against 12 months of historical data")
    add_callout_box(doc, "~$3.30/month total operating cost")
    add_callout_box(doc, "Zero new dependencies -- backtester uses only numpy")

    add_spacer(doc, 4)

    add_body_text(doc,
        "Target user: A buy-and-hold investor who also trades rhythmic and swing patterns "
        "and wants actionable options trade ideas with specific strikes, premiums, and risk profiles.",
        bold_terms=["Target user:"]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "2. System Architecture", level=1)

    add_body_text(doc,
        "The system comprises 13 modules orchestrated by a single entry point (main1.py), "
        "deployed as two Azure Functions -- a timer trigger for daily automation and an "
        "HTTP trigger for on-demand analysis. Continuous integration and deployment are "
        "handled by GitHub Actions."
    )

    add_section_heading(doc, "2.1 Component Overview", level=2)

    component_headers = ["Component", "Purpose", "Lines of Code"]
    component_rows = [
        ["main1.py", "Pipeline orchestrator -- loads tickers, runs modules, builds report", "~260"],
        ["technical.py", "22 technical indicators (EMA, RSI, MACD, Bollinger, VWAP, SMA)", "~185"],
        ["fundamentals.py", "~22 fundamental fields (P/E, growth, targets, short interest)", "~100"],
        ["options_monitor.py", "Options chain analysis, IV, max pain, unusual activity, skew", "~250"],
        ["options_strategy.py", "7-strategy engine with weighted scoring and risk profiles", "~500"],
        ["backtester.py", "Black-Scholes options strategy backtester (numpy only)", "~290"],
        ["backtester_entry_exit.py", "Entry/exit signal backtester for swing trade validation", "~150"],
        ["market_sentiment.py", "StockTwits social sentiment (bullish/bearish ratio, score)", "~75"],
        ["news.py", "Google News RSS feed with article content scraping", "~25"],
        ["strategy.py", "Entry/exit signal evaluator (RSI, VWAP, EMA crossover logic)", "~200"],
        ["summarizer.py", "GPT-5 structured summarization (300-word constrained output)", "~160"],
        ["report_builder.py", "HTML dashboard, strategy cards, catalyst calendar", "~420"],
        ["emailer.py", "Gmail SMTP delivery (SSL, port 465)", "~25"],
    ]
    create_styled_table(doc, component_headers, component_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "2.2 System Architecture Diagram", level=2)

    add_body_text(doc,
        "The following table illustrates the end-to-end data flow from trigger sources "
        "through the orchestrator, into analysis modules, and out to delivery channels."
    )

    # Architecture flow as styled table
    arch_headers = ["Layer", "Components", "Description"]
    arch_rows = [
        ["TRIGGERS", "Timer (M-F 1PM UTC)  |  HTTP API  |  CLI",
         "Three entry points invoke the pipeline orchestrator"],
        ["    -->", "", ""],
        ["ORCHESTRATOR", "main1.py",
         "Loads watchlist from Key Vault, iterates tickers, coordinates all modules"],
        ["    -->", "", ""],
        ["ANALYSIS MODULES", "fundamentals.py  ->  technical.py  ->  news.py  ->  strategy.py  ->  options_monitor.py  ->  options_strategy.py  ->  backtester.py  ->  market_sentiment.py  ->  summarizer.py",
         "Sequential per-ticker execution; each module fail-safe with graceful fallback"],
        ["    -->", "", ""],
        ["DATA SOURCES", "yfinance (free)  |  Google News RSS (free)  |  StockTwits API (free)  |  OpenAI GPT-5 (pay-per-token)",
         "Four external APIs; only OpenAI incurs cost"],
        ["    -->", "", ""],
        ["OUTPUTS", "HTML Email Dashboard  |  JSON API Response  |  Terminal Output",
         "Rich HTML report with dashboard table, detail cards, strategy cards, catalyst calendar"],
    ]
    create_styled_table(doc, arch_headers, arch_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "2.3 Design Principles", level=2)

    principles = [
        ("Fail-safe per module:", "Each module is wrapped in try/except at the orchestrator level. "
         "If options data fails, the pipeline continues with None and the report gracefully omits that section."),
        ("No new dependencies:", "Options analysis uses yfinance's built-in option_chain(). Sentiment uses "
         "the existing requests package. Zero additional pip packages were introduced."),
        ("Backward-compatible signatures:", "All modified function signatures use keyword arguments with "
         "defaults, so existing callers work unchanged."),
        ("Import-time safety:", "New modules are imported with try/except fallbacks to None, so the timer "
         "trigger continues to work even if a new module has an import error."),
    ]

    for title, desc in principles:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.left_indent = Cm(0.5)
        run_t = p.add_run(title + "  ")
        run_t.bold = True
        run_t.font.size = Pt(10.5)
        run_t.font.name = "Calibri"
        run_t.font.color.rgb = DARK_BLUE
        run_d = p.add_run(desc)
        run_d.font.size = Pt(10.5)
        run_d.font.name = "Calibri"
        run_d.font.color.rgb = DARK_GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 3. DATA PIPELINE
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "3. Data Pipeline", level=1)

    add_body_text(doc,
        "The pipeline follows a deliberate execution order designed to satisfy data dependencies "
        "between modules. Each step produces output that downstream modules consume."
    )

    add_section_heading(doc, "3.1 Pipeline Flow", level=2)

    pipeline_steps = [
        ["Trigger Fires", "Timer (weekday 1PM UTC), HTTP request, or CLI invocation",
         "Calls main1.run(tickers, email_flag, format)"],
        ["Load Watchlist", "Azure Key Vault, TICKERS env var, CLI args, or HTTP params",
         "Configurable list of equity tickers"],
        ["Fundamentals", "yfinance stock.info + calendar + earnings_dates",
         "Provides last_earnings_date used by technical module for VWAP anchoring"],
        ["Technicals", "yfinance 1-year daily OHLCV with 3-tier fetch fallback",
         "22 indicators; provides stock_price used by options module for ATM identification"],
        ["News + Strategy", "Google News RSS (5 items) + RSI/VWAP/EMA signal evaluation",
         "Independent modules with no cross-dependencies"],
        ["Options Analysis", "yfinance option_chain + 7-strategy scoring engine",
         "Needs stock price from technicals; produces IV, max pain, strategy recommendations"],
        ["Backtesting (opt.)", "Black-Scholes walk-forward simulation on 1-year history",
         "Validates strategy recommendations against historical data; activated via --backtest flag"],
        ["Sentiment", "StockTwits API (up to 30 messages with self-tagged sentiment)",
         "Independent; computes bullish/bearish ratio and sentiment score"],
        ["AI Summary", "OpenAI GPT-5 with structured prompt and constrained output",
         "Receives all upstream data; generates 300-word analysis per ticker"],
        ["Report Build", "HTML dashboard table + per-ticker detail + strategy cards + calendar",
         "Inline CSS for email client compatibility (Gmail, Outlook)"],
        ["Delivery", "Gmail SMTP (timer) or HTTP response (API) or terminal (CLI)",
         "Email includes full dashboard; API returns HTML or JSON"],
    ]

    create_flow_table(doc, pipeline_steps, title_col="Stage", desc_col="Data Source / Action",
                      detail_col="Key Detail")

    add_spacer(doc, 8)

    add_section_heading(doc, "3.2 Data Sources Comparison", level=2)

    ds_headers = ["Source", "Data Provided", "Cost", "Auth Required", "Reliability"]
    ds_rows = [
        ["yfinance (v0.2.65)", "OHLCV, .info (182 fields), options chains, splits, calendar",
         "Free", "None", "Good -- 3-tier fallback handles edge cases"],
        ["StockTwits API", "30 messages with user-tagged sentiment (Bullish/Bearish)",
         "Free", "None", "Moderate -- undocumented rate limits, currently 403"],
        ["Google News RSS", "Top 5 news items with titles and article content",
         "Free", "None", "Moderate -- paywalls, bot detection, 15-60 min lag"],
        ["OpenAI GPT-5", "300-word structured analysis per ticker",
         "~$0.01/1K tokens", "API Key", "High -- consistent structured output"],
    ]
    create_styled_table(doc, ds_headers, ds_rows)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 4. TECHNICAL ANALYSIS ENGINE
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "4. Technical Analysis Engine", level=1)

    add_body_text(doc,
        "The technical analysis module computes 22 indicators per ticker from one year of "
        "daily OHLCV (Open, High, Low, Close, Volume) data. These indicators span five categories: "
        "trend, momentum, volatility, volume, and price structure. Together, they provide a "
        "comprehensive view of each stock's current technical posture."
    )

    add_callout_box(doc, "22 technical indicators computed per ticker per run")

    add_spacer(doc, 4)

    add_section_heading(doc, "4.1 Technical Indicators by Category", level=2)

    ti_headers = ["Category", "Indicator", "Method / Description", "Signal Interpretation"]
    ti_rows = [
        ["Trend", "EMA(9), EMA(20)", "Exponential moving averages",
         "Short-term trend direction; crossovers signal momentum shifts"],
        ["Trend", "SMA(50)", "Simple moving average, 50-day window",
         "Medium-term trend; price above = bullish"],
        ["Trend", "SMA(200)", "Simple moving average, 200-day window (requires 1-year data)",
         "Long-term secular trend; price above = bull market"],
        ["Momentum", "RSI(14)", "Relative Strength Index via Wilder's smoothing",
         "<30 = oversold, >70 = overbought"],
        ["Momentum", "MACD Line", "EMA(12) - EMA(26)",
         "Positive = bullish momentum"],
        ["Momentum", "MACD Signal", "EMA(9) of MACD line",
         "MACD crossing above signal = buy signal"],
        ["Momentum", "MACD Histogram", "MACD line - signal line",
         "Positive histogram = bullish momentum"],
        ["Volatility", "Bollinger Upper", "SMA(20) + 2 x StdDev(20)",
         "Upper band; price touching = potential overbought"],
        ["Volatility", "Bollinger Lower", "SMA(20) - 2 x StdDev(20)",
         "Lower band; price touching = potential oversold"],
        ["Volatility", "BB Width", "(Upper - Lower) / SMA(20)",
         "<0.04 = Bollinger squeeze, low volatility preceding breakout"],
        ["Volume", "Volume Ratio", "Today's volume / 20-day average volume",
         ">2.0 = unusual activity, potential institutional interest"],
        ["Volume", "Anchored VWAP", "Volume-weighted avg price from last earnings date",
         "Price vs VWAP indicates fair value alignment"],
        ["Price Structure", "Support (20d)", "20-day rolling low",
         "Nearest demand zone / floor"],
        ["Price Structure", "Resistance (20d)", "20-day rolling high",
         "Nearest supply zone / ceiling"],
        ["Price Structure", "52-Week High/Low", "Max and min over 1 year",
         "Full range context for position sizing"],
        ["Price Structure", "% Changes", "1-day, 5-day, 1-month (21d), 3-month (63d)",
         "Multi-timeframe momentum snapshot"],
    ]
    create_styled_table(doc, ti_headers, ti_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "4.2 Data Fetching Resilience", level=2)

    add_body_text(doc,
        "The data fetching layer implements a 3-tier fallback strategy to handle yfinance's "
        "inconsistent behavior across different tickers and market conditions. All results pass "
        "through a column normalization function that handles MultiIndex columns and ensures "
        "consistent naming. This design ensures the pipeline rarely fails due to data retrieval issues.",
        bold_terms=["3-tier fallback strategy"]
    )

    fetch_headers = ["Tier", "Method", "When Used"]
    fetch_rows = [
        ["1 (Primary)", "yf.download(auto_adjust=True)", "Fastest path; works for most tickers"],
        ["2 (Fallback)", "Ticker.history(auto_adjust=True)", "Different code path; succeeds when download fails"],
        ["3 (Last Resort)", "yf.download(auto_adjust=False)", "Raw unadjusted data; covers remaining edge cases"],
    ]
    create_styled_table(doc, fetch_headers, fetch_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "4.3 VWAP Implementation", level=2)

    add_body_text(doc,
        "The Volume Weighted Average Price uses an anchored approach. The primary calculation "
        "anchors VWAP from the last earnings date, resetting each earnings cycle to stay relevant "
        "to current price action. If the earnings date is unavailable, the system falls back to a "
        "20-day rolling VWAP. This replaced an earlier cumulative VWAP implementation that produced "
        "divergent, nonsensical values over multi-month periods.",
        bold_terms=["anchored approach"]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 5. OPTIONS ANALYSIS & STRATEGY ENGINE
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "5. Options Analysis & Strategy Engine", level=1)

    add_body_text(doc,
        "The options analysis system is composed of two tightly integrated modules: "
        "options_monitor.py for raw chain analytics, and options_strategy.py for the "
        "7-strategy recommendation engine. Together, they transform raw options chain data "
        "into scored, actionable trade recommendations with computed risk profiles.",
        bold_terms=["options_monitor.py", "options_strategy.py"]
    )

    add_callout_box(doc, "7 algorithmically-scored options strategies with exact risk profiles")

    add_spacer(doc, 4)

    add_section_heading(doc, "5.1 Options Chain Analytics", level=2)

    add_body_text(doc,
        "The options monitor targets the nearest monthly expiry with 15 to 50 days to expiration "
        "(DTE), preferring approximately 30 DTE to capture meaningful time value while avoiding "
        "near-expiry theta decay. Key metrics computed include:"
    )

    opt_metrics = [
        ("ATM Implied Volatility:", "Average of at-the-money call and put IV from the options chain. "
         "Represents the market's forward-looking volatility expectation."),
        ("Max Pain:", "The strike price at which the total value of all in-the-money options is "
         "minimized. Market makers theoretically benefit when price settles here at expiration."),
        ("Put/Call Ratio:", "Computed from both volume and open interest. High ratios signal bearish "
         "positioning; low ratios signal bullish sentiment."),
        ("Unusual Activity:", "Flags strikes where volume exceeds 2x open interest, indicating new "
         "positions being opened -- a potential signal of informed directional bets."),
        ("IV Skew:", "OTM put IV minus OTM call IV (5% out-of-the-money). Positive skew indicates "
         "demand for downside protection; negative skew indicates speculative call buying."),
    ]

    for title, desc in opt_metrics:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(3)
        pf.left_indent = Cm(0.5)
        run_t = p.add_run(title + "  ")
        run_t.bold = True
        run_t.font.size = Pt(10.5)
        run_t.font.name = "Calibri"
        run_t.font.color.rgb = DARK_GRAY
        run_d = p.add_run(desc)
        run_d.font.size = Pt(10.5)
        run_d.font.name = "Calibri"
        run_d.font.color.rgb = DARK_GRAY

    add_spacer(doc, 6)

    add_section_heading(doc, "5.2 Strategy Scoring Matrix", level=2)

    add_body_text(doc,
        "Each of the seven strategies is evaluated against five weighted conditions drawn from "
        "technical indicators, implied volatility, and fundamental catalysts. The weighted score "
        "(0 to 1) determines the recommendation status: Recommended (score >= 0.60 with at least "
        "3 conditions met), Monitor (score >= 0.40), or Avoid."
    )

    strat_headers = ["#", "Strategy", "Market View", "Key Entry Conditions", "Weight Focus"]
    strat_rows = [
        ["1", "Covered Call", "Neutral-bullish income",
         "Price > SMA50, IV 25-45%, RSI 40-60", "Trend + IV level"],
        ["2", "Cash-Secured Put", "Bullish on pullback",
         "Price near support, IV > 30%, RSI < 40", "Support + oversold"],
        ["3", "Bull Call Spread", "Directional bullish",
         "EMA9 > EMA20, Price > VWAP, IV 20-50%", "Momentum + value"],
        ["4", "Bear Call Spread", "Directional bearish",
         "EMA9 < EMA20, RSI > 65, IV > 40%", "Trend reversal + IV"],
        ["5", "Iron Condor", "Range-bound / neutral",
         "BB width < 0.06, IV 40-70%, RSI 40-60", "Volatility squeeze"],
        ["6", "Protective Put", "Long stock hedge",
         "Price near resistance, earnings within DTE", "Event risk + price level"],
        ["7", "Long Straddle", "Big move expected",
         "Earnings 5-15 days, unusual activity, IV < 50%", "Catalyst + cheap IV"],
    ]
    create_styled_table(doc, strat_headers, strat_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "5.3 Options Strategy Risk Profiles", level=2)

    add_body_text(doc,
        "Each recommendation includes computed risk metrics based on actual chain premiums. "
        "The table below summarizes the theoretical risk profile for each strategy type."
    )

    risk_headers = ["Strategy", "Max Profit", "Max Loss", "Breakeven", "Risk/Reward Profile"]
    risk_rows = [
        ["Covered Call", "Premium received", "Stock decline - premium",
         "Stock price - premium", "Limited upside, downside stock risk offset by premium"],
        ["Cash-Secured Put", "Premium received", "Strike - premium (if assigned)",
         "Strike - premium", "Obligation to buy at strike; premium provides cushion"],
        ["Bull Call Spread", "Spread width - net debit", "Net debit paid",
         "Long strike + net debit", "Defined risk; capped profit and loss"],
        ["Bear Call Spread", "Net credit received", "Spread width - net credit",
         "Short strike + net credit", "Defined risk; profits from decline or flat market"],
        ["Iron Condor", "Net credit received", "Spread width - net credit",
         "Two breakevens (upper and lower)", "Profits from range-bound price action"],
        ["Protective Put", "Unlimited (less put cost)", "Put premium paid",
         "Stock price + put cost", "Insurance; protects existing long position"],
        ["Long Straddle", "Unlimited (both directions)", "Total premium paid",
         "Strike +/- total premium", "Profits from large move in either direction"],
    ]
    create_styled_table(doc, risk_headers, risk_rows)

    add_spacer(doc, 6)

    add_body_text(doc,
        "Strike selection uses the full options chain DataFrames to find optimal strikes: ATM strikes "
        "via nearest-to-price selection, OTM strikes via percentage-based targeting (e.g., 5% OTM "
        "for protective puts), and spread widths based on price-relative ranges (3-10% for bull call spreads)."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 6. BLACK-SCHOLES BACKTESTING FRAMEWORK
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "6. Black-Scholes Backtesting Framework", level=1)

    add_body_text(doc,
        "The backtesting module validates options strategy recommendations against historical data "
        "using Black-Scholes pricing. Rather than relying on historical options chain data (which "
        "is not freely available), the system synthesizes option prices from historical stock prices "
        "and volatility, enabling walk-forward testing over a full year of data.",
        bold_terms=["Black-Scholes pricing"]
    )

    add_callout_box(doc, "Black-Scholes backtesting validates strategies against 12 months of historical data")
    add_callout_box(doc, "Zero new dependencies -- backtester uses only numpy")

    add_spacer(doc, 4)

    add_section_heading(doc, "6.1 Backtest Methodology", level=2)

    bt_steps = [
        ["Fetch Historical Data", "Download 1-year daily OHLCV data using the same 3-tier fallback as the live pipeline"],
        ["Compute Daily Indicators", "Calculate all technical indicators (RSI, EMA, MACD, Bollinger, VWAP, historical volatility) for each trading day"],
        ["Check Entry Conditions", "Apply the same condition-checking logic used by the live strategy engine to identify entry points"],
        ["Price Strategy Legs (BS)", "On entry signal, price each option leg using the Black-Scholes formula with historical volatility as the IV proxy"],
        ["Simulate Holding Period", "Hold position for the target DTE or until an exit signal triggers"],
        ["Reprice at Exit (BS)", "Compute exit P&L using Black-Scholes repricing at the simulated exit date"],
        ["Aggregate Results", "Compile win rate, average return, max drawdown, and profit factor across all simulated trades"],
    ]
    create_flow_table(doc, bt_steps, title_col="Step", desc_col="Description")

    add_spacer(doc, 8)

    add_section_heading(doc, "6.2 Black-Scholes Implementation", level=2)

    add_body_text(doc,
        "The Black-Scholes formula is implemented using numpy's error function for the "
        "normal cumulative distribution -- providing exact computation without requiring scipy. "
        "The key formula components are:"
    )

    bs_headers = ["Variable", "Formula", "Description"]
    bs_rows = [
        ["d1", "[ ln(S/K) + (r + v^2/2)T ] / (v * sqrt(T))", "Primary distribution parameter"],
        ["d2", "d1 - v * sqrt(T)", "Secondary distribution parameter"],
        ["Call Price", "S * N(d1) - K * e^(-rT) * N(d2)", "European call option value"],
        ["Put Price", "K * e^(-rT) * N(-d2) - S * N(-d1)", "European put option value"],
    ]
    create_styled_table(doc, bs_headers, bs_rows)

    add_body_text(doc,
        "Where S = stock price, K = strike price, r = risk-free rate, v = volatility, "
        "T = time to expiration (years), and N() = cumulative normal distribution function."
    )

    add_spacer(doc, 4)

    add_section_heading(doc, "6.3 Accuracy and Limitations", level=2)

    add_body_text(doc,
        "Black-Scholes pricing estimates are approximately 70-80% accurate compared to real market "
        "premiums. Key limitations include the constant volatility assumption (real IV fluctuates), "
        "no bid-ask spread modeling (real markets have friction costs), and European-style exercise "
        "only (most US equity options are American-style with early exercise rights). Despite these "
        "limitations, the backtester provides a statistically useful validation of strategy logic."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 7. AI SUMMARIZATION (GPT-5)
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "7. AI Summarization (GPT-5)", level=1)

    add_body_text(doc,
        "The AI summarization module is the final per-ticker processing step. It receives all "
        "upstream data -- technicals, fundamentals, options, sentiment, strategy signals, and news -- "
        "and produces a structured 300-word analysis using OpenAI's GPT-5 model. The system prompt "
        "establishes a persona of an experienced derivatives trader with dual analytical lenses.",
        bold_terms=["GPT-5"]
    )

    add_section_heading(doc, "7.1 Structured Output Format", level=2)

    add_body_text(doc,
        "The AI output is constrained to exactly six sections, ensuring consistency across "
        "tickers and runs:"
    )

    output_headers = ["Section", "Content", "Audience Value"]
    output_rows = [
        ["Verdict", "One-line directional call (Bullish, Bearish, Neutral)",
         "Instant decision framework"],
        ["Buy & Hold Lens", "2-3 bullets on long-term thesis",
         "Relevant for retirement / core portfolio positions"],
        ["Swing/Rhythmic Setup", "Entry, stop, and target prices",
         "Actionable trade with defined risk parameters"],
        ["Options Play", "Specific strategy, strike, expiry, premium",
         "Uses pre-computed strategy engine recommendation"],
        ["What Changed Today", "1-2 bullets on new developments",
         "Delta from previous session -- what is different now"],
        ["Risk Flag", "One-line risk warning",
         "Forces acknowledgment of key risk before acting"],
    ]
    create_styled_table(doc, output_headers, output_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "7.2 Token Economics", level=2)

    add_body_text(doc,
        "Input is structured as labeled sections (TECHNICALS, FUNDAMENTALS, OPTIONS, SENTIMENT, "
        "STRATEGY, NEWS) rather than raw Python dictionaries, reducing token usage by approximately "
        "60% compared to the previous design that sent full article content."
    )

    token_headers = ["Component", "Estimated Tokens"]
    token_rows = [
        ["System prompt", "~200"],
        ["Technicals section", "~200-300"],
        ["Fundamentals section", "~150-200"],
        ["Options section", "~100-150"],
        ["Sentiment section", "~80-120"],
        ["Strategy section", "~50-80"],
        ["News titles (5 items)", "~100-150"],
        ["Total Input", "~900-1,200"],
        ["Output (300 words)", "~400-600"],
        ["Total per Ticker", "~1,300-1,800"],
    ]
    create_styled_table(doc, token_headers, token_rows)

    add_spacer(doc, 6)

    add_body_text(doc,
        "A key design decision is that the AI receives the pre-computed strategy recommendation "
        "from the options strategy engine rather than improvising its own options play. This ensures "
        "the Options Play section in the AI summary is grounded in algorithmically-validated conditions "
        "and real chain premiums, not hallucinated strikes or premiums."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 8. REPORT GENERATION
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "8. Report Generation", level=1)

    add_body_text(doc,
        "The report builder produces a single HTML document designed for email delivery. All "
        "styling uses inline CSS because email clients (Gmail, Outlook) strip style tags and "
        "external stylesheets. The report is structured for quick scanning -- a dashboard table "
        "provides a single-glance overview, followed by detailed cards for each ticker.",
        bold_terms=["inline CSS"]
    )

    add_section_heading(doc, "8.1 Report Components", level=2)

    report_headers = ["Component", "Content", "Styling"]
    report_rows = [
        ["Dashboard Table", "Ticker, Price, 1D%, Signal, RSI, ATM IV, Key Level, Sentiment, Verdict",
         "Color-coded badges: green=BUY, red=EXIT, yellow=OVERBOUGHT, blue=OVERSOLD"],
        ["Per-Ticker Detail", "GPT summary (markdown to HTML), key metrics, technicals, fundamentals",
         "Card layout with bordered sections"],
        ["Options Snapshot Bar", "Expiry, DTE, ATM IV, call/put premiums, P/C ratio, max pain",
         "Compact single-line bar within detail card"],
        ["Strategy Cards", "Recommended strategies with status, score, strike/premium, risk profile",
         "Color-coded: green=Recommended, yellow=Monitor, gray=Avoid"],
        ["Backtest Results Card", "Win rate, avg return, max drawdown, profit factor, trade count",
         "Summary statistics with pass/fail indicators"],
        ["Catalyst Calendar", "Upcoming earnings and ex-dividend dates across all tickers",
         "Sorted by proximity; date-formatted table"],
        ["Compact Technicals", "SMA50, SMA200, VWAP, MACD histogram, BB width, volume ratio",
         "Single line per ticker in detail section"],
        ["Compact Fundamentals", "Market cap, P/E, recommendation, target price, short interest",
         "Single line per ticker in detail section"],
    ]
    create_styled_table(doc, report_headers, report_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "8.2 Dashboard Table Columns", level=2)

    add_body_text(doc,
        "The dashboard opens the report with a single-glance summary across all tickers. Each column "
        "is designed to answer a specific question a portfolio manager would ask:"
    )

    dash_headers = ["Column", "Source Module", "Question Answered"]
    dash_rows = [
        ["Ticker", "--", "What am I looking at?"],
        ["Price", "technical.py", "Where is it trading right now?"],
        ["1D%", "technical.py", "How did it move today?"],
        ["Signal", "strategy.py", "Is there a BUY, EXIT, or NEUTRAL signal?"],
        ["RSI", "technical.py", "Is it overbought or oversold?"],
        ["ATM IV", "options_monitor.py", "How expensive are options right now?"],
        ["Key Level", "technical.py", "What is the nearest support or resistance?"],
        ["Sentiment", "market_sentiment.py", "What is the social crowd saying?"],
        ["Verdict", "summarizer.py", "What is the AI's one-line conclusion?"],
    ]
    create_styled_table(doc, dash_headers, dash_rows)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 9. CLOUD INFRASTRUCTURE (AZURE)
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "9. Cloud Infrastructure (Azure)", level=1)

    add_body_text(doc,
        "The system is deployed on Microsoft Azure using the serverless Azure Functions platform. "
        "This eliminates server management overhead and provides automatic scaling, with costs "
        "driven entirely by execution count rather than uptime.",
        bold_terms=["Azure Functions"]
    )

    add_section_heading(doc, "9.1 Azure Resource Map", level=2)

    azure_headers = ["Resource", "Name / ID", "Purpose"]
    azure_rows = [
        ["Function App", "stock-daily-runner", "Hosts both timer and HTTP trigger functions"],
        ["Resource Group", "rg-stocks", "Logical container for all resources"],
        ["Region", "Central US", "Deployment location"],
        ["Key Vault", "stockdailyvault20172025", "Stores API keys, ticker list, and secrets"],
        ["Storage Account", "dailystockstorage", "ZIP deployment packages, function state"],
        ["App Insights", "45c7a0c4-...", "Monitoring, logging, exception tracking"],
        ["App Service Plan", "ASP-rgstocks-8c6b", "Consumption plan (pay-per-execution)"],
    ]
    create_styled_table(doc, azure_headers, azure_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "9.2 Function Triggers", level=2)

    trigger_headers = ["Aspect", "DailyRunner (Timer)", "StockAnalysisHttp (HTTP)"]
    trigger_rows = [
        ["Trigger Type", "Timer: Mon-Fri 1PM UTC", "HTTP: GET or POST"],
        ["Authentication", "N/A (internal Azure trigger)", "Function-level API key"],
        ["Email Delivery", "Yes (always sends)", "No (returns report in response)"],
        ["Ticker Source", "Azure Key Vault", "Request parameters (query or body)"],
        ["Output", "Side effect: email sent", "HTML or JSON response"],
        ["Use Case", "Daily automated intelligence", "Ad-hoc analysis on demand"],
    ]
    create_styled_table(doc, trigger_headers, trigger_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "9.3 Deployment Model", level=2)

    add_body_text(doc,
        "The function app runs with WEBSITE_RUN_FROM_PACKAGE=1, which mounts the deployment "
        "ZIP as a read-only filesystem. This provides faster cold starts (no pip install on startup), "
        "deterministic deployments (exact same package every time), and a smaller attack surface "
        "(no server-side build process). The Oryx build system is explicitly disabled.",
        bold_terms=["WEBSITE_RUN_FROM_PACKAGE=1"]
    )

    add_section_heading(doc, "9.4 CI/CD Pipeline", level=2)

    add_body_text(doc,
        "GitHub Actions automates testing and deployment on every push to main. The pipeline "
        "runs two sequential jobs:"
    )

    cicd_headers = ["Job", "Steps", "Purpose"]
    cicd_rows = [
        ["Test", "1. Validate requirements.txt (pinned versions)\n2. Install runtime deps\n3. Install test deps\n4. Run pytest",
         "Ensures code quality and dependency integrity before deployment"],
        ["Deploy", "1. Prebuild .python_packages (3 retries)\n2. Create azure namespace init\n3. Verify vendored packages\n4. Create release.zip\n5. Verify zip contents (7 checks)\n6. Azure login (OIDC)\n7. Set run-from-package=1\n8. Disable Oryx build\n9. Deploy via Azure Functions action\n10. Post-deploy: verify 2 functions listed",
         "Packages and deploys with comprehensive pre- and post-deploy verification"],
    ]
    create_styled_table(doc, cicd_headers, cicd_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "9.5 Security Model", level=2)

    sec_headers = ["Secret", "Storage Location", "Access Method"]
    sec_rows = [
        ["OPENAI_API_KEY", "Azure Key Vault", "Function App Setting (Key Vault reference)"],
        ["EMAIL_USER", "Azure App Settings", "Environment variable"],
        ["EMAIL_PASS", "Azure App Settings", "Environment variable"],
        ["KEY_VAULT_NAME", "Azure App Settings", "Environment variable"],
        ["Tickers Watchlist", "Azure Key Vault", "SecretClient with DefaultAzureCredential (Managed Identity)"],
        ["HTTP Function Key", "Azure Functions", "Query parameter (?code=) or x-functions-key header"],
    ]
    create_styled_table(doc, sec_headers, sec_rows)

    add_body_text(doc,
        "The function app authenticates to Key Vault using Managed Identity via "
        "DefaultAzureCredential. No credentials for Key Vault itself are stored in code or "
        "environment variables. Local development uses a .env file excluded from git."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 10. COST ANALYSIS
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "10. Cost Analysis", level=1)

    add_body_text(doc,
        "One of the system's most compelling attributes is its extremely low operating cost. "
        "By leveraging Azure's free tier for compute and storage, free data APIs for market "
        "data and sentiment, and efficient prompt engineering for AI summarization, the total "
        "monthly cost is approximately $3.30."
    )

    add_callout_box(doc, "~$3.30/month total operating cost")

    add_spacer(doc, 4)

    add_section_heading(doc, "10.1 Monthly Cost Breakdown (3 Tickers, Weekday Runs)", level=2)

    cost_headers = ["Service", "Unit Cost", "Monthly Usage", "Monthly Cost"]
    cost_rows = [
        ["Azure Functions", "Free tier: 1M executions/mo", "~22 runs x 1 execution", "$0.00"],
        ["Azure Key Vault", "$0.03 per 10K operations", "~66 operations/mo", "$0.00"],
        ["Azure Storage", "$0.0184/GB/mo", "<1 MB", "$0.00"],
        ["Application Insights", "Free tier: 5 GB/mo", "<100 MB", "$0.00"],
        ["OpenAI GPT-5", "~$0.01/1K tokens", "~5K tokens x 3 tickers x 22 days = 330K tokens", "~$3.30"],
        ["yfinance", "Free (scrapes Yahoo Finance)", "--", "$0.00"],
        ["StockTwits API", "Free (no auth required)", "--", "$0.00"],
        ["Google News RSS", "Free", "--", "$0.00"],
    ]
    create_styled_table(doc, cost_headers, cost_rows)

    add_spacer(doc, 4)

    # Total row as a callout
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("TOTAL MONTHLY COST:  ~$3.30")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK_BLUE

    add_spacer(doc, 4)

    add_body_text(doc,
        "OpenAI API usage is the only recurring cost. All Azure services operate within free tier "
        "limits, and all market data sources are free. The cost scales linearly with the number of "
        "tickers: approximately $1.10 per ticker per month. Adding a fourth ticker would bring the "
        "total to roughly $4.40 per month."
    )

    add_spacer(doc, 4)

    add_section_heading(doc, "10.2 Cost Scaling Projection", level=2)

    scale_headers = ["Tickers", "Daily Token Usage", "Monthly Token Usage", "Monthly Cost"]
    scale_rows = [
        ["3 (current)", "~15K tokens/day", "~330K tokens/mo", "~$3.30"],
        ["5", "~25K tokens/day", "~550K tokens/mo", "~$5.50"],
        ["10", "~50K tokens/day", "~1.1M tokens/mo", "~$11.00"],
        ["20", "~100K tokens/day", "~2.2M tokens/mo", "~$22.00"],
    ]
    create_styled_table(doc, scale_headers, scale_rows)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 11. ROADMAP
    # ═══════════════════════════════════════════════════════════

    add_section_heading(doc, "11. Roadmap", level=1)

    add_body_text(doc,
        "The system has undergone two major enhancement cycles since initial deployment. "
        "The February 2026 release (v3.0) added the options strategy engine, Black-Scholes "
        "backtesting, and strategy cards in the report. The following tables summarize completed "
        "work and planned enhancements."
    )

    add_section_heading(doc, "11.1 Completed Enhancements (v3.0)", level=2)

    done_headers = ["Enhancement", "Date", "Description"]
    done_rows = [
        ["Options strategy engine", "2026-02-09",
         "7 defined-risk strategies with weighted scoring and computed risk profiles"],
        ["Black-Scholes backtesting", "2026-02-09",
         "Walk-forward BS-simulated backtesting for all 7 options strategies"],
        ["Entry/exit signal backtesting", "2026-02-09",
         "Historical validation of RSI/VWAP/EMA crossover swing trade signals"],
        ["Daily IV persistence", "2026-02-09",
         "Auto-appends ATM IV to iv_history.csv each run for future IV Rank computation"],
        ["Strategy cards in report", "2026-02-09",
         "Color-coded strategy recommendation cards with risk profiles in the email report"],
        ["GPT strategy integration", "2026-02-09",
         "GPT uses the pre-computed strategy recommendation instead of improvising options plays"],
        ["Options chain analysis", "2026-02-08",
         "Full chain analytics: ATM IV, max pain, unusual activity, IV skew, P/C ratio"],
        ["Market sentiment (StockTwits)", "2026-02-08",
         "Replaced Reddit/PRAW dependency with zero-auth StockTwits API"],
        ["Anchored VWAP", "2026-02-08",
         "Fixed broken cumulative VWAP with earnings-anchored calculation"],
        ["HTML dashboard report", "2026-02-08",
         "Dashboard table + detail cards + catalyst calendar (replaced text blocks)"],
        ["HTTP trigger + CLI", "2026-02-08",
         "Added HTTP API and CLI access modes alongside existing timer trigger"],
    ]
    create_styled_table(doc, done_headers, done_rows)

    add_spacer(doc, 8)

    add_section_heading(doc, "11.2 Future Roadmap", level=2)

    future_headers = ["Enhancement", "Priority", "Complexity", "Expected Benefit"]
    future_rows = [
        ["IV Rank computation", "High", "Low",
         "Contextualize current IV against 52-week historical range (requires ~252 daily snapshots)"],
        ["Alternative sentiment source", "High", "Medium",
         "Replace currently non-functional StockTwits API (403 errors)"],
        ["Multi-timeframe analysis", "Medium", "Low",
         "Confirm daily signals with weekly and monthly timeframe trends"],
        ["Azure Blob IV persistence", "Medium", "Medium",
         "Durable IV history storage that survives deployments (currently local CSV)"],
        ["Bollinger squeeze alert", "Low", "Low",
         "Proactively flag imminent breakout setups when BB width contracts"],
        ["Mobile push notifications", "Low", "Medium",
         "Faster alert delivery than email for time-sensitive signals"],
        ["Remove legacy praw/nltk deps", "Low", "Low",
         "Reduce package size and deploy time by removing unused Reddit dependencies"],
    ]
    create_styled_table(doc, future_headers, future_rows)

    add_spacer(doc, 12)

    # ═══════════════════════════════════════════════════════════
    # FOOTER / CLOSING
    # ═══════════════════════════════════════════════════════════

    # Decorative line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("_" * 60)
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10)

    # Closing info
    for text in [
        "Document generated: February 2026",
        "Pipeline version: 3.0 (options strategy engine + backtesting)",
        "Author: cdemchalk",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = MEDIUM_GRAY

    # ─── Add Page Numbers ───────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Daily Stock Analysis Pipeline  |  Page X" to footer
        run1 = p.add_run("Daily Stock Analysis Pipeline   |   Page ")
        run1.font.size = Pt(8)
        run1.font.name = "Calibri"
        run1.font.color.rgb = MEDIUM_GRAY

        # Page number field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_field = p.add_run()
        run_field._r.append(fld_char_begin)

        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_instr = p.add_run()
        run_instr._r.append(instr_text)
        run_instr.font.size = Pt(8)
        run_instr.font.name = "Calibri"
        run_instr.font.color.rgb = MEDIUM_GRAY

        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_end = p.add_run()
        run_end._r.append(fld_char_end)

    # ─── Save Document ──────────────────────────────────────────
    output_path = "/mnt/e/github/daily_stock_analysis/docs/Daily_Stock_Analysis_Whitepaper.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Whitepaper saved to: {output_path}")
    print(f"File size: {os.path.getsize(output_path):,} bytes")


if __name__ == "__main__":
    generate_whitepaper()
